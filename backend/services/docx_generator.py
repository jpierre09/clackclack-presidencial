"""Generate reclamation DOCX documents from the departamental template."""

from __future__ import annotations

import io
import zipfile
from datetime import date, datetime

from docx import Document

from backend.config import FORMATO_DOCX

_MONTHS_ES = {
    1: "Enero",
    2: "Febrero",
    3: "Marzo",
    4: "Abril",
    5: "Mayo",
    6: "Junio",
    7: "Julio",
    8: "Agosto",
    9: "Septiembre",
    10: "Octubre",
    11: "Noviembre",
    12: "Diciembre",
}
_NAME_PLACEHOLDER = "________________________________________________________________________________"
_CC_PLACEHOLDER = "________________________________________"


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph while keeping the first run formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
        return
    paragraph.text = new_text


def _build_date_text(reference_date: date | datetime | None = None) -> str:
    today = reference_date or datetime.now()
    return f"{today.day} {_MONTHS_ES[today.month]} de {today.year}"


def _build_claimant_text(user_name: str = "", user_cc: str = "") -> str:
    claimant_name = user_name.strip() or _NAME_PLACEHOLDER
    claimant_cc = user_cc.strip() or _CC_PLACEHOLDER
    return (
        f"Yo, {claimant_name}, con c\u00e9dula de ciudadan\u00eda {claimant_cc} mayor de edad "
        "identificado (a) c\u00f3mo aparece al pie de mi firma, en mi condici\u00f3n de testigo electoral "
        "debidamente acreditado (a) por la organizaci\u00f3n electoral, a trav\u00e9s de la presente, "
        "de conformidad con lo expuesto en los art\u00edculos 164 y 192 del C\u00f3digo Electoral "
        "( Decreto 2241 de 1986), solicit\u00f3 de forma respetuosa RECUENTO DE VOTOS para la elecci\u00f3n "
        "de PRESIDENTE Y VICEPRESIDENTE, la cual sustent\u00f3 en los siguientes:"
    )


def _build_location_text(alert_data: dict) -> str:
    municipio = str(alert_data.get("municipio") or alert_data.get("municipio_cod") or "").strip()
    zona = str(alert_data.get("zona_cod") or "").zfill(2)
    puesto = str(alert_data.get("puesto_nombre") or alert_data.get("puesto_cod") or "").strip()
    mesa = int(alert_data.get("mesa") or 0)
    pres_votes = alert_data.get("pres_votes_actual")
    pres_text = "N/A" if pres_votes is None else str(pres_votes)

    return (
        f"En el municipio de {municipio}, Zona {zona}, Puesto {puesto}, "
        f"Mesa {mesa}, se evidencia una anomal\u00eda en los votos "
        f"registrados para la elecci\u00f3n presidencial: Total votos f\u00f3rmulas: {pres_text}."
    )


def _build_grouped_location_text(alerts: list[dict]) -> str:
    if not alerts:
        return ""

    municipio = str(alerts[0].get("municipio") or alerts[0].get("municipio_cod") or "").strip()
    ordered_alerts = sorted(
        alerts,
        key=lambda alert: (
            str(alert.get("zona_cod") or "").zfill(2),
            str(alert.get("puesto_cod") or "").zfill(2),
            int(alert.get("mesa") or 0),
        ),
    )

    lines = [
        (
            f"En el municipio de {municipio}, se evidencian las siguientes anomal\u00edas "
            "en los votos registrados para la elecci\u00f3n de PRESIDENTE Y VICEPRESIDENTE:"
        )
    ]
    for idx, alert_data in enumerate(ordered_alerts, start=1):
        zona = str(alert_data.get("zona_cod") or "").zfill(2)
        puesto_cod = str(alert_data.get("puesto_cod") or "").zfill(2)
        puesto = str(alert_data.get("puesto_nombre") or puesto_cod).strip()
        mesa = int(alert_data.get("mesa") or 0)
        pres_votes = alert_data.get("pres_votes_actual")
        pres_text = "N/A" if pres_votes is None else str(pres_votes)
        lines.append(
            f"{idx}. Zona {zona}, Puesto {puesto} (c\u00f3digo {puesto_cod}), Mesa {mesa}: "
            f"Total votos f\u00f3rmulas presidenciales: {pres_text}."
        )

    return "\n".join(lines)


def _apply_common_template_fields(
    doc: Document,
    fundamentos_text: str,
    user_name: str = "",
    user_cc: str = "",
    reference_date: date | datetime | None = None,
) -> None:
    date_text = _build_date_text(reference_date)
    claimant_text = _build_claimant_text(user_name, user_cc)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if text.startswith("Yo, "):
            _set_paragraph_text(paragraph, claimant_text)
            continue

        if text.startswith("En el municipio de "):
            _set_paragraph_text(paragraph, fundamentos_text)
            continue

        if any(month in text for month in _MONTHS_ES.values()) and "de 20" in text:
            _set_paragraph_text(paragraph, date_text)


def generate_single(
    alert_data: dict,
    comision: dict | None = None,
    user_name: str = "",
    user_cc: str = "",
    reference_date: date | datetime | None = None,
) -> io.BytesIO:
    """Generate a single DOCX using the departamental template."""
    del comision

    doc = Document(str(FORMATO_DOCX))
    location_text = _build_location_text(alert_data)
    _apply_common_template_fields(doc, location_text, user_name, user_cc, reference_date)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_grouped_by_municipio(
    alerts: list[dict],
    user_name: str = "",
    user_cc: str = "",
    reference_date: date | datetime | None = None,
) -> io.BytesIO:
    """Generate a single DOCX containing all alert fundamentos for one municipality."""
    if not alerts:
        raise ValueError("Se requiere al menos una alerta para generar la reclamacion consolidada.")

    doc = Document(str(FORMATO_DOCX))
    fundamentos_text = _build_grouped_location_text(alerts)
    _apply_common_template_fields(doc, fundamentos_text, user_name, user_cc, reference_date)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


async def generate_batch(
    alerts: list[dict],
    user_name: str = "",
    user_cc: str = "",
    reference_date: date | datetime | None = None,
) -> io.BytesIO:
    """Generate a ZIP containing one DOCX per alert."""
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for alert in alerts:
            docx_buffer = generate_single(alert, None, user_name, user_cc, reference_date)
            municipio = str(alert.get("municipio") or alert.get("municipio_cod") or "MUN").replace(" ", "_")
            filename = (
                f"{municipio}/"
                f"Reclamacion_Departamental_Z{str(alert['zona_cod']).zfill(2)}_"
                f"P{str(alert['puesto_cod']).zfill(2)}_M{int(alert['mesa']):03d}.docx"
            )
            zf.writestr(filename, docx_buffer.read())

    zip_buffer.seek(0)
    return zip_buffer
